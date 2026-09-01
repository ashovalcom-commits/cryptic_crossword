import re
import docx
from grid import Grid

def extract_text_from_docx(file_path: str) -> str:
    """
    קוראת קובץ docx, מזהה משבצות שחורות ותאים ממוזגים בטבלה.
    """
    doc = docx.Document(file_path)
    extracted_lines = []
    
    # 1. חילוץ הלוח (הטבלאות)
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            seen_cells = set() # למניעת כפילויות במקרה של תאים ממוזגים
            
            for cell in row.cells:
                if cell in seen_cells:
                    continue
                seen_cells.add(cell)
                
                xml_str = cell._tc.xml
                
                # בדיקה לכמה משבצות התא הזה נמתח (מיזוג תאים)
                span = 1
                span_match = re.search(r'<w:gridSpan w:val="(\d+)"', xml_str)
                if span_match:
                    span = int(span_match.group(1))
                    
                # בדיקה האם יש צבע רקע (משבצת שחורה)
                is_black = False
                if 'w:fill=' in xml_str:
                    fill_match = re.search(r'w:fill="([^"]+)"', xml_str)
                    # מוודאים שהצבע הוא לא הגדרת "ללא צבע" או "לבן"
                    if fill_match and fill_match.group(1).lower() not in ['auto', 'ffffff', 'clear']:
                        is_black = True
                        
                if is_black:
                    # הוספת סולמיות לפי גודל המיזוג
                    row_data.append('#' * span)
                else:
                    clean_cell = cell.text.replace('\n', '').strip()
                    row_data.append(clean_cell if clean_cell else " ")
                    
            # סוגרים כל שורה עם מפריד בסוף כדי שנדע מתי היא נחתכת
            extracted_lines.append("|".join(row_data) + "|")
            
    # נוסיף סימן ברור שמפריד בין הלוח לחלק של ההגדרות הטקסטואליות
    extracted_lines.append("---BOARD_END---")
    
    # 2. חילוץ ההגדרות (הפסקאות)
    for para in doc.paragraphs:
        if para.text.strip():
            extracted_lines.append(para.text.strip())
            
    return '\n'.join(extracted_lines)


def parse_grid_to_matrix(raw_text: str, cols: int = 11) -> list:
    """
    מקבלת את הטקסט המעובד ומחזירה מטריצה של נקודות, סולמיות ומספרי עוגנים.
    """
    board_section = raw_text.split("---BOARD_END---")[0].strip()
    processed_text = board_section.replace('\n', '')
    raw_cells = [c for c in processed_text.split('|') if c != '']
    
    cells_list = []
    for cell in raw_cells:
        clean_cell = cell.strip()
        if '#' in clean_cell:
            cells_list.extend(['#'] * clean_cell.count('#'))
        elif clean_cell.isdigit():
            # עכשיו אנחנו שומרים את המספר במקום למחוק אותו!
            cells_list.append(clean_cell)
        else:
            cells_list.append('.')
            
    matrix = []
    for i in range(0, len(cells_list), cols):
        row = cells_list[i:i+cols]
        if len(row) == cols:
            matrix.append(row)
            
    return matrix


def parse_clues(raw_text: str) -> dict:
    """
    מחלצת את טקסט ההגדרות ומחלקת אותן לאופקי ואנכי.
    מתמודדת עם הגדרות מפוצלות (כמו "23+8 אופקי") וקרדיטים.
    """
    clues_section = raw_text.split("---BOARD_END---")[1]
    
    if "אנכי:" in clues_section:
        across_part, down_part = clues_section.split("אנכי:")
        across_part = across_part.replace("אופקי:", "")
    else:
        across_part = clues_section.replace("אופקי:", "")
        down_part = ""

    parsed_clues = {"ACROSS": {}, "DOWN": {}}
    
    # תבנית לזיהוי תחילת הגדרה: מספר, אופציונלי +מספר, אופציונלי "אופקי/אנכי", ונקודה.
    clue_start_pattern = re.compile(r'^(\d+(?:\+\d+)?(?:\s*אופקי|\s*אנכי)?)\.', re.MULTILINE)
    
    def extract_clues_from_text(text_block, direction_dict):
        matches = list(clue_start_pattern.finditer(text_block))
        for i, match in enumerate(matches):
            clue_id = match.group(1).strip()
            
            # חותכים מהנקודה של ההגדרה הנוכחית ועד תחילת ההגדרה הבאה
            start_idx = match.end()
            end_idx = matches[i+1].start() if i + 1 < len(matches) else len(text_block)
            
            raw_clue_text = text_block[start_idx:end_idx].strip()
            # מאחדים שורות שגלשו
            raw_clue_text = " ".join(raw_clue_text.split())
            
            # מחפשים את תבנית האורך בכל מקום בתוך הטקסט (כמו 3,4 או 7)
            length_match = re.search(r'\(([\d\,]+)\)', raw_clue_text)
            lengths = length_match.group(1) if length_match else ""
            
            if length_match:
                # מוציאים את תבנית האורך מהמקום שבו נמצאה ומצמידים אותה לסוף ההגדרה,
                # כדי שמספר האותיות יופיע בסוף ולא בתחילת הטקסט
                raw_clue_text = (
                    raw_clue_text[:length_match.start()] + raw_clue_text[length_match.end():]
                ).strip()
                raw_clue_text = " ".join(raw_clue_text.split())
                raw_clue_text = f"{raw_clue_text} ({lengths})"
            
            direction_dict[clue_id] = {
                "text": raw_clue_text,
                "lengths": lengths
            }

    extract_clues_from_text(across_part, parsed_clues["ACROSS"])
    extract_clues_from_text(down_part, parsed_clues["DOWN"])

    return parsed_clues

def link_clues_to_grid(grid: Grid, parsed_clues: dict):
    """
    מזריקה את טקסט ההגדרות לתוך העוגנים המתאימים בלוח.
    """
    for slot in grid.slots:
        if not slot.clue_number:
            continue
            
        dir_key = "ACROSS" if slot.direction.name == "ACROSS" else "DOWN"
        clues_dict = parsed_clues[dir_key]
        
        # מחפשים את ההגדרה במילון
        for key, data in clues_dict.items():
            # שימוש בביטוי רגולרי כדי ש-'23' יתפוס גם את '23' וגם את '23+8 אופקי'
            if re.match(rf"^{slot.clue_number}(?:\+|$|\s)", str(key)):
                slot.clue_text = data["text"]
                break


if __name__ == "__main__":
    file_path = r"C:\Users\ashov\Downloads\28-08-26.docx"
    
    try:
        # 1. קריאת המסמך הגולמי
        raw_text = extract_text_from_docx(file_path)
        
        # 2. פענוח גיאומטריה והגדרות
        result_matrix = parse_grid_to_matrix(raw_text)
        parsed_clues = parse_clues(raw_text)
        
        # 3. יצירת הלוח החכם (אורכי מילים מינימליים בתשבץ כזה הם 2)
        my_grid = Grid(result_matrix, min_word_length=2)
        
        # 4. שידוך ההגדרות לעוגנים
        link_clues_to_grid(my_grid, parsed_clues)
        
        # 5. הדפסת בקרה - בוא נראה אם זה עבד!
        print("--- הצגת עוגנים נבחרים עם ההגדרות שלהם ---")
        for slot in my_grid.slots:
            if slot.clue_number:
                print(f"[{slot.clue_number} {slot.direction.name}] אורך: {slot.length} | הגדרה: {slot.clue_text}")
                
    except Exception as e:
        print(f"שגיאה בתהליך: {e}")